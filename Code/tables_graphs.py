# tables_graphs.py - Python implementation of tables_graphs.do.
# Tables (missing="-"; area share = "-" where the attribute is unmeasured) + plot/household trend figures.
#!/usr/bin/env python3
# Build the 3 weighted share tables + faceted-by-country trend graphs from
# rental_tenure_ALL.dta. Design-based (Taylor-linearized) 95% CIs with
# strata = country|wave|strataid, PSU = country|wave|ea_id, single-PSU strata
# centered (contribute 0), matching Stata svyset ... singleunit(centered).
import pyreadstat, numpy as np, pandas as pd, os
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D

import os as _os
_cands=[
 "/Users/jchamberlin/Library/CloudStorage/Dropbox/SSA-pooled-survey-data/ssa-land-rental",
 "/Users/jchamberlin/Library/CloudStorage/Dropbox/LSMS-ISA-harmonised-dataset-on-agricultural-productivity-and-welfare/ssa-land-rental",
]
ROOT=next((c for c in _cands if _os.path.exists(c+"/Output/Final/rental_tenure_ALL.dta")), _cands[-1])
FINAL=f"{ROOT}/Output/Final"; TAB=f"{ROOT}/Output/Tables"; FIG=f"{ROOT}/Output/Figures"
os.makedirs(TAB,exist_ok=True); os.makedirs(FIG,exist_ok=True)

IND=["parcel_rentedin","parcel_rentedout","parcel_purchased","parcel_certificate"]
NICE={"parcel_rentedin":"Rented-in","parcel_rentedout":"Rented-out",
      "parcel_purchased":"Purchased","parcel_certificate":"Has certificate"}
CORDER=["Ethiopia","Malawi","Mali","Niger","Nigeria","Tanzania","Uganda","Zambia"]

df,_=pyreadstat.read_dta(f"{FINAL}/rental_tenure_ALL.dta")
df=df[df["season"]==1].copy()                       # report season 1
for v in IND: df[v]=pd.to_numeric(df[v],errors="coerce")
df["weight"]=pd.to_numeric(df["weight"],errors="coerce")
df["parcel_area_ha"]=pd.to_numeric(df["parcel_area_ha"],errors="coerce")
df["strataid"]=df["strataid"].astype(str); df["ea_id"]=df["ea_id"].astype(str)
df=df[df["weight"].notna() & (df["weight"]>0)]
df["_strat"]=df["country"]+"|"+df["wave"].astype(str)+"|"+df["strataid"]
df["_psu"]  =df["country"]+"|"+df["wave"].astype(str)+"|"+df["ea_id"]

def lin_var(e, strat, psu):
    """Taylor-linearized variance of a total of linearized values e,
    stratified design; single-PSU strata contribute 0 (centered)."""
    d=pd.DataFrame({"e":e,"h":strat,"i":psu})
    u=d.groupby(["h","i"])["e"].sum()                # PSU totals
    var=0.0
    for h,grp in u.groupby(level=0):
        nh=len(grp)
        if nh<2: continue                            # singleton stratum -> 0
        var+= nh/(nh-1.0)*((grp-grp.mean())**2).sum()
    return var

def svymean(sub,y):
    """Weighted mean of binary y with design-based 95% CI; returns (est,lo,hi,n)."""
    s=sub[sub[y].notna()]
    if len(s)==0 or s[y].notna().sum()==0: return (np.nan,np.nan,np.nan,0)
    w=s["weight"].values; yy=s[y].values
    W=w.sum(); est=(w*yy).sum()/W
    e=w*(yy-est)/W
    se=np.sqrt(lin_var(e,s["_strat"].values,s["_psu"].values))
    return (est, est-1.96*se, est+1.96*se, len(s))

def svyratio_area(sub,y):
    """Area-weighted share = sum(w*area*y)/sum(w*area), design CI."""
    s=sub[sub[y].notna() & sub["parcel_area_ha"].notna() & (sub["parcel_area_ha"]>0)]
    if len(s)==0: return (np.nan,np.nan,np.nan,0)
    # If the attribute IS observed in the cell but never among area-measured parcels
    # (e.g. rented-out parcels are uncultivated and carry no measured area), the area
    # share is undefined, not zero -> return missing.
    if (sub[y]==1).sum()>0 and (s[y]==1).sum()==0:
        return (np.nan,np.nan,np.nan,0)
    w=s["weight"].values; a=s["parcel_area_ha"].values; yy=s[y].values
    D=(w*a).sum(); est=(w*a*yy).sum()/D
    e=w*a*(yy-est)/D
    se=np.sqrt(lin_var(e,s["_strat"].values,s["_psu"].values))
    return (est, est-1.96*se, est+1.96*se, len(s))

# household-level frame: 1 row per (country,wave,hh); indicator = max over season-1 plots
def hh_frame(d):
    keys=["country","wave","year","hh_id"]
    agg={v:"max" for v in IND}
    agg.update({"weight":"first","_strat":"first","_psu":"first"})
    h=d.groupby(keys,as_index=False).agg(agg)
    return h
HH=hh_frame(df)

rows=[]
for c in CORDER:
    for yr in sorted(df.loc[df.country==c,"year"].unique()):
        cell_p=df[(df.country==c)&(df.year==yr)]
        cell_h=HH[(HH.country==c)&(HH.year==yr)]
        for v in IND:
            ep=svymean(cell_p,v); eh=svymean(cell_h,v); ea=svyratio_area(cell_p,v)
            rows.append(dict(country=c,year=int(yr),indicator=v,
                             hh=eh[0],hh_lo=eh[1],hh_hi=eh[2],
                             plot=ep[0],plot_lo=ep[1],plot_hi=ep[2],
                             area=ea[0],area_lo=ea[1],area_hi=ea[2],
                             n_plot=ep[3],n_hh=eh[3]))
R=pd.DataFrame(rows)
R.to_csv(f"{TAB}/shares_long_with_CIs.csv",index=False)

# ---- wide tables (rows country-year, cols 4 indicators) per level ----
def wide(level):
    t=R.pivot_table(index=["country","year"],columns="indicator",values=level,sort=False)
    t=t.reindex(columns=IND); t.columns=[NICE[c] for c in IND]
    t=t.reset_index()
    t["country"]=pd.Categorical(t["country"],categories=CORDER,ordered=True)
    return t.sort_values(["country","year"]).reset_index(drop=True)
T={lvl:wide(lvl) for lvl in ["hh","plot","area"]}
TITLES={"hh":"Share of households with one or more plot",
        "plot":"Share of plots","area":"Share of farm area (ha)"}
# Presentation view: country labelled once per group (blank on repeat rows),
# values to 3 decimals, structural-missing cells blank. (The long file
# shares_long_with_CIs.csv keeps full labels + CIs for machine use.)
def present(t):
    t=t.copy()
    t["country"]=t["country"].astype(str)          # drop Categorical so "" is allowed
    for col in [NICE[i] for i in IND]:
        t[col]=t[col].map(lambda x:"-" if pd.isna(x) else f"{x:.3f}")   # "-" = not asked that round
    t["year"]=t["year"].astype(int).astype(str)
    dup=t["country"].eq(t["country"].shift())
    t.loc[dup,"country"]=""
    t=t.rename(columns={"country":"Country","year":"Year"})
    return t[["Country","Year"]+[NICE[i] for i in IND]]

for lvl,t in T.items():
    present(t).to_csv(f"{TAB}/table_{lvl}_share.csv",index=False)

# Excel: one sheet per table
with pd.ExcelWriter(f"{TAB}/tenure_share_tables.xlsx") as xl:
    for lvl,t in T.items():
        present(t).to_excel(xl,sheet_name=TITLES[lvl][:31],index=False)

# Word
from docx import Document
doc=Document()
doc.add_heading("Land-tenure & rental-market descriptives",0)
doc.add_paragraph("Survey-weighted shares by country and survey year (season 1). "
    "Blank cells = the question was not asked that round (structural missing; see provenance). "
    "Cross-country levels should be read with care: survey instruments and panel structure "
    "differ across countries and waves.")
for lvl in ["hh","plot","area"]:
    p=present(T[lvl]); doc.add_heading(TITLES[lvl],level=1)
    tbl=doc.add_table(rows=1,cols=len(p.columns)); tbl.style="Light Grid Accent 1"
    for j,cn in enumerate(p.columns): tbl.rows[0].cells[j].text=cn
    for _,r in p.iterrows():
        cs=tbl.add_row().cells
        for j,cn in enumerate(p.columns):
            cs[j].text=str(r[cn])
            if j==0 and r[cn]!="":              # bold the country label
                cs[j].paragraphs[0].runs[0].font.bold=True
doc.save(f"{TAB}/tenure_share_tables.docx")

# ---- faceted-by-country trend graphs, plot-level, with CI bands ----
COL={"parcel_rentedin":"#1f77b4","parcel_rentedout":"#d62728",
     "parcel_purchased":"#2ca02c","parcel_certificate":"#9467bd"}
fig,axes=plt.subplots(2,4,figsize=(16,8),sharey=False)   # free y per panel (each 0-based)
for ax,c in zip(axes.ravel(),CORDER):
    sub=R[R.country==c]
    yrs=sorted(sub.year.unique())
    pmax=np.nanmax(sub[["plot_hi","plot"]].values)
    for v in IND:
        s=sub[sub.indicator==v].sort_values("year")
        ax.plot(s["year"],s["plot"],marker="o",ms=4,color=COL[v],lw=1.6)
        ax.fill_between(s["year"],s["plot_lo"],s["plot_hi"],color=COL[v],alpha=0.15,linewidth=0)
    ax.set_title(c,fontsize=11,fontweight="bold")
    ax.set_xticks(yrs); ax.set_xticklabels(yrs,rotation=45,fontsize=8)
    ax.set_ylim(0, (pmax*1.15 if np.isfinite(pmax) and pmax>0 else 1)); ax.grid(alpha=0.25)
axes[0,0].set_ylabel("share of plots"); axes[1,0].set_ylabel("share of plots")
handles=[Line2D([0],[0],color=COL[v],marker="o",lw=1.6,label=NICE[v]) for v in IND]
fig.legend(handles=handles,loc="lower center",ncol=4,frameon=False,fontsize=11,bbox_to_anchor=(0.5,-0.02))
fig.suptitle("Plot-level tenure & rental-market shares over time, by country (95% CI bands; note: y-axis scaled per country)",
             fontsize=13,fontweight="bold")
fig.tight_layout(rect=[0,0.03,1,0.97])
fig.savefig(f"{FIG}/trends_by_country_plot.png",dpi=150,bbox_inches="tight")

# ---- faceted-by-country HOUSEHOLD-level trend lines with CI bands (same style as plot) ----
figh,axesh=plt.subplots(2,4,figsize=(16,8),sharey=False)
for ax,c in zip(axesh.ravel(),CORDER):
    sub=R[R.country==c]
    yrs=sorted(sub.year.unique())
    pmax=np.nanmax(sub[["hh_hi","hh"]].values)
    for v in IND:
        s=sub[sub.indicator==v].sort_values("year")
        ax.plot(s["year"],s["hh"],marker="o",ms=4,color=COL[v],lw=1.6)
        ax.fill_between(s["year"],s["hh_lo"],s["hh_hi"],color=COL[v],alpha=0.15,linewidth=0)
    ax.set_title(c,fontsize=11,fontweight="bold")
    ax.set_xticks(yrs); ax.set_xticklabels(yrs,rotation=45,fontsize=8)
    ax.set_ylim(0,(pmax*1.15 if np.isfinite(pmax) and pmax>0 else 1)); ax.grid(alpha=0.25)
axesh[0,0].set_ylabel("share of households"); axesh[1,0].set_ylabel("share of households")
handlesh=[Line2D([0],[0],color=COL[v],marker="o",lw=1.6,label=NICE[v]) for v in IND]
figh.legend(handles=handlesh,loc="lower center",ncol=4,frameon=False,fontsize=11,bbox_to_anchor=(0.5,-0.02))
figh.suptitle("Household-level shares over time, by country (share of households with >=1 plot; 95% CI bands; y-axis scaled per country)",
              fontsize=13,fontweight="bold")
figh.tight_layout(rect=[0,0.03,1,0.97])
figh.savefig(f"{FIG}/trends_by_country_hh.png",dpi=150,bbox_inches="tight")

print("rows in long table:",len(R))
print("wrote tables to",TAB)
print("wrote figure to",FIG)
# quick console peek of plot-level table
pd.set_option("display.width",160,"display.max_columns",20)
print("\nPLOT-LEVEL share (head):")
print(T["plot"].head(12).to_string(index=False))
